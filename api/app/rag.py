import os
import re
import uuid
from typing import List, Optional, Dict, Tuple

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

USE_OPENAI = os.getenv("USE_OPENAI", "1") == "1"

from langchain_core.prompts import ChatPromptTemplate

from .schemas import (
    ComplianceVerdict,
    ContractChangeResponse,
    ContractChangeResult,
    ContractExtract,
    EmailDraftContent,
    EmailDraft,
    PolicyRuleExtract,
)

# -----------------------------
# LLM + Embeddings
# -----------------------------

def build_embeddings():
    if USE_OPENAI:
        from langchain_openai import OpenAIEmbeddings
        return OpenAIEmbeddings(model=os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small"))

    if os.getenv("USE_OLLAMA_EMBEDDINGS", "0") == "1":
        from langchain_ollama import OllamaEmbeddings

        return OllamaEmbeddings(
            model=os.getenv("OLLAMA_EMBED_MODEL", "nomic-embed-text"),
            base_url=os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"),
        )

    try:
        from langchain_huggingface import HuggingFaceEmbeddings
    except ImportError:
        from langchain_community.embeddings import HuggingFaceEmbeddings

    return HuggingFaceEmbeddings(model_name=os.getenv(
        "HF_EMBED_MODEL", "sentence-transformers/all-MiniLM-L6-v2"
    ))

def build_llm():
    if USE_OPENAI:
        from langchain_openai import ChatOpenAI

        return ChatOpenAI(
            model=os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini"),
            temperature=0,
        )

    from langchain_ollama import ChatOllama

    return ChatOllama(
        model=os.getenv("OLLAMA_CHAT_MODEL", "qwen2.5:7b"),
        temperature=0,
        base_url=os.getenv("OLLAMA_BASE_URL", "http://localhost:11434"),
    )


def _infer_question_scope(question: str) -> Optional[str]:
    normalized = (question or "").lower()

    contract_markers = [
        "договор",
        "сотрудник",
        "работник",
        "контракт",
        "email",
        "фио",
    ]
    if any(token in normalized for token in contract_markers):
        return "contract"

    policy_markers = [
        "политик",
        "регламент",
        "правил",
        "заявлен",
        "отпуск",
        "отпуска",
        "части отпуска",
        "дней до начала отпуска",
    ]
    if any(token in normalized for token in policy_markers):
        return "policy"

    return None


def _extract_doc_ids_from_question(question: str) -> List[str]:
    if not question:
        return []
    found = re.findall(r"\b(?:policy|contract)_[a-zA-Z0-9]+\b", question, flags=re.IGNORECASE)
    deduped: List[str] = []
    seen = set()
    for item in found:
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


def _normalize_unit(unit: Optional[str]) -> Optional[str]:
    if unit is None:
        return None

    normalized = unit.strip().lower()
    if not normalized:
        return None

    aliases = {
        "дней": "календарных дней",
        "дня": "календарных дней",
        "день": "календарных дней",
        "calendar days": "календарных дней",
        "calendars days": "календарных дней",
        "day": "календарных дней",
        "days": "календарных дней",
    }
    return aliases.get(normalized, normalized)


def _extract_number(text: Optional[str]) -> Optional[int]:
    if not text:
        return None

    prioritized_patterns = [
        r"(\d+)\s+(?:календарных\s+)?дн(?:ей|я|ь)\b",
        r"(\d+)\s+calendar\s+days\b",
        r"(\d+)\s+days\b",
        r"not\s+less\s+than\s+(\d+)\b",
        r"не\s+менее\s+(\d+)\b",
    ]
    for pattern in prioritized_patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return int(match.group(1))

    numbers = re.findall(r"(\d+)", text)
    if not numbers:
        return None

    # Часто первое число — номер пункта вроде 4.3, а нужное число стоит дальше.
    return int(numbers[-1])


def _extract_unit(text: Optional[str]) -> Optional[str]:
    if not text:
        return None

    patterns = [
        r"\d+\s+([A-Za-z]+(?:\s+[A-Za-z]+)?)",
        r"\d+\s+([А-Яа-яЁё]+(?:\s+[А-Яа-яЁё]+)?)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            unit = match.group(1).strip(" .,:;)")
            if unit:
                return unit
    return None


def _extract_emails(text: Optional[str]) -> List[str]:
    if not text:
        return []
    found = re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    deduped: List[str] = []
    seen = set()
    for email in found:
        lower = email.lower()
        if lower in seen:
            continue
        seen.add(lower)
        deduped.append(email)
    return deduped


def _clean_draft_body(text: str) -> str:
    body = text.strip()
    body = re.sub(r"^Body:\s*", "", body, flags=re.IGNORECASE)
    body = re.sub(r"^\*\*Тело письма:\*\*\s*", "", body, flags=re.IGNORECASE)
    body = re.sub(r"^\*\*Body:\*\*\s*", "", body, flags=re.IGNORECASE)
    body = re.sub(r"\[Your Name\].*", "", body, flags=re.IGNORECASE | re.S)
    body = re.sub(r"\[Your Position\].*", "", body, flags=re.IGNORECASE | re.S)
    body = re.sub(r"\[Company Name\].*", "", body, flags=re.IGNORECASE | re.S)
    body = re.sub(r"\[Ваше имя\].*", "", body, flags=re.IGNORECASE | re.S)
    body = re.sub(r"\[Название организации\].*", "", body, flags=re.IGNORECASE | re.S)
    body = re.sub(r"Please find attached.*", "", body, flags=re.IGNORECASE)
    body = re.sub(r"review the attached document.*", "", body, flags=re.IGNORECASE)
    body = re.sub(r"kindly provide us with a supporting document.*", "", body, flags=re.IGNORECASE)
    body = re.sub(r"Пожалуйста, уведомите нас.*", "", body, flags=re.IGNORECASE)
    body = re.sub(r"Для подтверждения изменений рекомендуется обратиться.*", "", body, flags=re.IGNORECASE)
    body = re.sub(r"\*\*(.*?)\*\*", r"\1", body)
    body = re.sub(r"\[ФИО сотрудника\]", "Коллега", body, flags=re.IGNORECASE)
    body = re.sub(r"С уважением,\s*$", "", body, flags=re.IGNORECASE)
    body = re.sub(r"\n{3,}", "\n\n", body)
    return body.strip()


def _mask_email(email: Optional[str]) -> str:
    if not email:
        return "(no email)"
    local, _, domain = email.partition("@")
    if not domain:
        return "***"
    visible = local[:1] if local else ""
    return f"{visible}***@{domain}"


def _mask_name(name: Optional[str]) -> str:
    if not name:
        return "(unknown)"
    return name[:1] + "***"


def _enrich_policy_rule(rule: PolicyRuleExtract, context: str) -> PolicyRuleExtract:
    if rule.new_value is None:
        rule.new_value = _extract_number(rule.source_quote) or _extract_number(context)

    if rule.unit is None:
        rule.unit = _extract_unit(rule.source_quote) or _extract_unit(context)

    return rule


def _enrich_contract_extract(contract: ContractExtract, context: str) -> ContractExtract:
    emails = contract.emails or []
    if not emails:
        emails = _extract_emails(contract.source_quote) or _extract_emails(context)
    contract.emails = emails

    if contract.email is None and emails:
        contract.email = emails[0]

    if contract.current_value is None:
        contract.current_value = (
            _extract_number(contract.source_quote)
            or _extract_number(contract.unit)
            or _extract_number(context)
        )

    if contract.unit is None:
        contract.unit = (
            _extract_unit(contract.source_quote)
            or _extract_unit(context)
        )
    else:
        contract.unit = _extract_unit(contract.unit) or contract.unit

    if contract.evidence is None:
        contract.evidence = contract.source_quote

    return contract

# -----------------------------
# Простые лоадеры (txt/pdf/docx)
# -----------------------------
def load_text_from_file(filename: str, content: bytes) -> str:
    name = filename.lower()
    if name.endswith(".txt") or name.endswith(".md"):
        return content.decode("utf-8", errors="ignore")

    if name.endswith(".pdf"):
        from pypdf import PdfReader
        import io
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for p in reader.pages:
            pages.append(p.extract_text() or "")
        return "\n".join(pages).strip()

    if name.endswith(".docx"):
        from docx import Document as DocxDocument
        import io
        f = io.BytesIO(content)
        doc = DocxDocument(f)
        return "\n".join([p.text for p in doc.paragraphs]).strip()

    # fallback
    return content.decode("utf-8", errors="ignore")

# -----------------------------
# Хранилище документов + RAG индекс
# -----------------------------
class RAGStore:
    def __init__(self):
        self.embeddings = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=int(os.getenv("CHUNK_SIZE", "900")),
            chunk_overlap=int(os.getenv("CHUNK_OVERLAP", "150")),
        )
        self.vectorstore: Optional[FAISS] = None

        # doc_id -> (doc_type, raw_text)
        self.docs: Dict[str, Tuple[str, str]] = {}

    def _ensure_embeddings(self):
        if self.embeddings is None:
            self.embeddings = build_embeddings()

    def load_documents(self, documents: List[Dict[str, str]], rebuild_index: bool = True) -> int:
        self.docs = {
            item["doc_id"]: (item["doc_type"], item["text"])
            for item in documents
        }
        if rebuild_index:
            return self._rebuild_index()
        self.vectorstore = None
        return len(self.docs)

    def upsert_text_document(self, *, doc_type: str, filename: str, text: str, doc_id: Optional[str] = None) -> Tuple[str, int]:
        if not text.strip():
            raise ValueError("Документ пустой после извлечения текста.")
        if doc_id is None:
            doc_id = f"{doc_type}_{uuid.uuid4().hex[:10]}"
        self.docs[doc_id] = (doc_type, text)
        return doc_id, self._rebuild_index()

    def _rebuild_index(self) -> int:
        all_docs: List[Document] = []
        for doc_id, (doc_type, text) in self.docs.items():
            all_docs.append(Document(
                page_content=text,
                metadata={"doc_id": doc_id, "doc_type": doc_type},
            ))

        if not all_docs:
            self.vectorstore = None
            return 0

        self._ensure_embeddings()
        chunks = self.text_splitter.split_documents(all_docs)
        self.vectorstore = FAISS.from_documents(chunks, self.embeddings)
        return len(chunks)

    def upsert_document(self, *, doc_type: str, filename: str, content: bytes, doc_id: Optional[str] = None) -> Tuple[str, int]:
        text = load_text_from_file(filename, content)
        if not text.strip():
            raise ValueError("Файл прочитан, но текст пустой (возможно, скан/картинки без текста).")

        if doc_id is None:
            doc_id = f"{doc_type}_{uuid.uuid4().hex[:10]}"

        self.docs[doc_id] = (doc_type, text)
        chunks_indexed = self._rebuild_index()
        return doc_id, chunks_indexed

    def list_docs(self):
        return [{"doc_id": k, "doc_type": v[0], "chars": len(v[1])} for k, v in self.docs.items()]

    def has_doc(self, doc_id: str, doc_type: Optional[str] = None) -> bool:
        if doc_id not in self.docs:
            return False
        if doc_type is None:
            return True
        return self.docs[doc_id][0] == doc_type

    def _search(self, query: str, k: int = 8, doc_ids: Optional[List[str]] = None, doc_type: Optional[str] = None) -> List[Document]:
        if self.vectorstore is None:
            return []

        docs = self.vectorstore.similarity_search(query, k=max(k * 3, 20))

        def ok(d: Document) -> bool:
            md = d.metadata or {}
            if doc_ids is not None and md.get("doc_id") not in set(doc_ids):
                return False
            if doc_type is not None and md.get("doc_type") != doc_type:
                return False
            return True

        filtered = [d for d in docs if ok(d)]
        return filtered[:k]

    def _build_context(self, docs: List[Document]) -> str:
        return "\n\n---\n\n".join(
            [f"[{d.metadata.get('doc_type')}:{d.metadata.get('doc_id')}]\n{d.page_content}" for d in docs]
        )

    def ask(self, question: str, doc_ids: Optional[List[str]] = None) -> Tuple[str, List[str]]:
        llm = build_llm()

        explicit_doc_ids = _extract_doc_ids_from_question(question)
        if explicit_doc_ids:
            doc_ids = explicit_doc_ids

        scope = _infer_question_scope(question)
        ctx_docs = self._search(question, k=6, doc_ids=doc_ids, doc_type=scope)
        if not ctx_docs and scope is not None:
            ctx_docs = self._search(question, k=6, doc_ids=doc_ids)
        context = self._build_context(ctx_docs)

        if not context.strip():
            return ("В документах нет ответа на этот вопрос.", [])

        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "Ты строгий помощник по кадровым документам.\n"
             "Правила ответа:\n"
             "1. Используй только факты из контекста.\n"
             "2. Не выдумывай и не достраивай выводы.\n"
             "3. Если вопрос задан на русском, отвечай только на русском.\n"
             "4. Не переключайся на английский, китайский или другой язык без прямой просьбы.\n"
             "5. Если ответа нет в контексте, верни ровно фразу: 'В документах нет ответа на этот вопрос.'\n"
             "6. Отвечай коротко и по существу, без лишних вступлений.\n"
             "7. Если вопрос про конкретный policy или contract, не обобщай на другие документы."),
            ("human",
             "Контекст:\n{context}\n\nВопрос:\n{question}\n\nОтвет:"),
        ])

        msg = prompt.format_messages(context=context, question=question)
        out = llm.invoke(msg)

        citations = []
        for d in ctx_docs[:3]:
            md = d.metadata or {}
            citations.append(f"{md.get('doc_type')}:{md.get('doc_id')}")

        return (getattr(out, "content", str(out))), sorted(set(citations))

    def check_compliance(self, doc_ids: Optional[List[str]] = None, focus: str = "отпуска") -> ComplianceVerdict:
        llm = build_llm()

        # отдельно достаём policy + документы
        policy_docs = self._search(f"политика компании {focus} правила требования", k=10, doc_ids=doc_ids, doc_type="policy")
        other_docs  = self._search(f"{focus} отпуск минимальная длительность перенос оплата график", k=14, doc_ids=doc_ids)

        policy_ctx = "\n\n---\n\n".join([f"[policy:{d.metadata.get('doc_id')}]\n{d.page_content}" for d in policy_docs])
        docs_ctx   = "\n\n---\n\n".join([f"[{d.metadata.get('doc_type')}:{d.metadata.get('doc_id')}]\n{d.page_content}" for d in other_docs])

        verdict_chain = llm.with_structured_output(ComplianceVerdict)

        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "Ты HR-compliance аналитик. Твоя задача: проверить соответствие предоставленных документов политике компании по отпускам.\n"
             "Правила:\n"
             "- Не выдумывай факты.\n"
             "- Если в документах нет нужной информации — явно укажи это в summary и recommendations.\n"
             "- В evidence приводи короткие цитаты (1-3 строки) и помечай источник [policy:ID] или [contract:ID].\n"),
            ("human",
             "Политика (фрагменты):\n{policy_ctx}\n\n"
             "Документы (фрагменты):\n{docs_ctx}\n\n"
             "Проверь соответствие по фокусу: {focus}\n"
             "Верни структурированный вердикт."),
        ])

        msg = prompt.format_messages(policy_ctx=policy_ctx, docs_ctx=docs_ctx, focus=focus)
        return verdict_chain.invoke(msg) # type: ignore

    def extract_policy_rule(self, policy_doc_id: str) -> PolicyRuleExtract:
        if not self.has_doc(policy_doc_id, "policy"):
            raise ValueError(f"Policy документ не найден: {policy_doc_id}")

        llm = build_llm()
        ctx_docs = self._search(
            "минимальная часть отпуска минимальная длительность части отпуска правило срок длительность отпуска",
            k=6,
            doc_ids=[policy_doc_id],
            doc_type="policy",
        )
        if not ctx_docs:
            raise ValueError(f"Не удалось найти контекст policy: {policy_doc_id}")

        chain = llm.with_structured_output(PolicyRuleExtract)
        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "Ты извлекаешь правило из кадровой policy.\n"
             "Работай строго по контексту.\n"
             "Не выдумывай и не интерпретируй лишнее.\n"
             "Верни только структурированные поля.\n"
             "Если числа нет, new_value=null.\n"
             "source_quote дай короткой точной цитатой из контекста."),
            ("human",
             "Контекст policy:\n{context}\n\n"
             "Извлеки правило про минимальную длительность части отпуска или ближайшее числовое правило по отпуску."),
        ])

        msg = prompt.format_messages(context=self._build_context(ctx_docs))
        result = chain.invoke(msg)
        result = _enrich_policy_rule(result, self._build_context(ctx_docs))
        if not result.rule_topic.strip():
            result.rule_topic = "минимальная длительность части отпуска"
        return result

    def extract_contract_data(self, contract_doc_id: str, rule_topic: str) -> ContractExtract:
        if not self.has_doc(contract_doc_id, "contract"):
            raise ValueError(f"Договор не найден: {contract_doc_id}")

        llm = build_llm()
        ctx_docs = self._search(
            f"{rule_topic} отпуск vacation annual paid vacation minimum duration email employee работник",
            k=6,
            doc_ids=[contract_doc_id],
            doc_type="contract",
        )
        if not ctx_docs:
            raise ValueError(f"Не удалось найти контекст договора: {contract_doc_id}")

        chain = llm.with_structured_output(ContractExtract)
        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "Ты извлекаешь данные из трудового договора.\n"
             "Работай строго по контексту.\n"
             "Верни только структурированные данные.\n"
             "Не выдумывай.\n"
             "Если email или число не найдены, верни null.\n"
             "source_quote дай короткой цитатой только из договора."),
            ("human",
             "Контекст договора:\n{context}\n\n"
             "Извлеки:\n"
             "- employee_name\n"
             "- email\n"
             "- current_value по теме '{rule_topic}'\n"
             "- unit\n"
             "- source_quote"),
        ])

        msg = prompt.format_messages(
            context=self._build_context(ctx_docs),
            rule_topic=rule_topic,
        )
        result = chain.invoke(msg)
        return _enrich_contract_extract(result, self._build_context(ctx_docs))

    def compare_rule_to_contract(
        self,
        doc_id: str,
        rule: PolicyRuleExtract,
        contract: ContractExtract,
    ) -> ContractChangeResult:
        if rule.new_value is None:
            raise ValueError("Не удалось извлечь числовое значение правила из policy.")

        rule_unit = _normalize_unit(rule.unit)
        contract_unit = _normalize_unit(contract.unit)

        result = ContractChangeResult(
            doc_id=doc_id,
            employee_name=contract.employee_name,
            email=contract.email,
            emails=contract.emails,
            old_value=contract.current_value,
            new_value=rule.new_value,
            unit=contract.unit or rule.unit,
            needs_change=False,
            reason="",
            source_quote=contract.source_quote,
            evidence=contract.evidence,
        )

        if contract.current_value is None:
            result.needs_change = False
            result.reason = (
                "В договоре не найдено явное условие о минимальной длительности отпуска "
                "(нужна ручная проверка)."
            )
            return result

        if rule_unit is not None and contract_unit is not None and rule_unit != contract_unit:
            result.needs_change = True
            result.reason = "Единицы измерения не совпадают, нужна ручная проверка."
            return result

        if contract.current_value < rule.new_value:
            result.needs_change = True
            result.reason = (
                f"Условие договора ниже policy: {contract.current_value} < {rule.new_value}."
            )
            return result

        result.needs_change = False
        result.reason = (
            f"Договор соответствует policy: {contract.current_value} >= {rule.new_value}."
        )
        return result

    def build_notification_draft(
        self,
        rule: PolicyRuleExtract,
        contract_result: ContractChangeResult,
    ) -> Tuple[str, str]:
        llm = build_llm()
        draft_chain = llm.with_structured_output(EmailDraftContent)
        prompt = ChatPromptTemplate.from_messages([
            ("system",
             "Ты помощник HR. Сформируй короткое и корректное уведомление.\n"
             "Не давай юридических трактовок, не делай обещаний.\n"
             "Тон: нейтральный, деловой.\n"
             "Язык ответа: русский.\n"
             "Не упоминай, что ты ИИ.\n"
             "Не выдумывай вложения, подписи, согласования, подтверждения, приложения, ссылки, документы или факты, которых нет во входе.\n"
             "Не используй placeholders вроде [Your Name].\n"
             "Не проси сотрудника прислать документы, подтверждения или дополнительные сведения.\n"
             "Не указывай сроки, если их нет во входе.\n"
             "Пиши коротко: тема + 2-3 абзаца.\n"
             "Верни только два поля structured output: subject и body."),
            ("human",
             "Сформируй письмо сотруднику о необходимости обновить условие договора.\n\n"
             "Данные:\n"
             "- ФИО: {employee_name}\n"
             "- Doc ID: {doc_id}\n"
             "- Текущее условие (минимальный отпуск): {current_min_days}\n"
             "- Новое требование (политика): минимум {required_min_days} дней\n"
             "- Подтверждающий фрагмент (evidence): {evidence}"),
        ])

        msgs = prompt.format_messages(
            employee_name=contract_result.employee_name or "Коллега",
            doc_id=contract_result.doc_id,
            current_min_days=contract_result.old_value,
            required_min_days=contract_result.new_value,
            evidence=contract_result.evidence or "(не указан)",
        )
        draft = draft_chain.invoke(msgs)
        subject = (draft.subject or "").strip() or "Уточнение условий трудового договора"
        body = (draft.body or "").strip() or "Черновик не сформирован."
        body = _clean_draft_body(body)
        return subject, body

    def collect_email_drafts(self, response: ContractChangeResponse) -> List[EmailDraft]:
        drafts: List[EmailDraft] = []
        for item in response.results:
            if not item.needs_change:
                continue
            if not item.draft_subject or not item.draft_body:
                continue
            drafts.append(EmailDraft(
                doc_id=item.doc_id,
                to=item.emails,
                subject=item.draft_subject,
                body=item.draft_body,
            ))
        return drafts

    def send_email_stub(self, drafts: List[EmailDraft]) -> List[EmailDraft]:
        for draft in drafts:
            print("=== EMAIL DRAFT ===")
            print("Doc ID:", draft.doc_id)
            print("To:", ", ".join([_mask_email(email) for email in draft.to]) if draft.to else "(no emails)")
            print("Subject:", draft.subject)
            print(draft.body)
            print()
        return drafts

    def analyze_contract_changes(
        self,
        policy_doc_id: str,
        contract_doc_ids: List[str],
    ) -> ContractChangeResponse:
        rule = self.extract_policy_rule(policy_doc_id)
        results: List[ContractChangeResult] = []

        for contract_doc_id in contract_doc_ids:
            try:
                contract = self.extract_contract_data(contract_doc_id, rule.rule_topic)
                result = self.compare_rule_to_contract(contract_doc_id, rule, contract)
            except Exception as exc:
                result = ContractChangeResult(
                    doc_id=contract_doc_id,
                    employee_name=None,
                    email=None,
                    emails=[],
                    old_value=None,
                    new_value=rule.new_value,
                    unit=rule.unit,
                    needs_change=False,
                    reason=f"Не удалось корректно обработать договор: {exc}",
                    source_quote=None,
                    evidence=None,
                    draft_subject=None,
                    draft_body=None,
                )

            if result.needs_change:
                subject, body = self.build_notification_draft(rule, result)
                result.draft_subject = subject
                result.draft_body = body

            results.append(result)

        return ContractChangeResponse(policy_rule=rule, results=results)
